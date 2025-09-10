"""
Document Processing Utilities for AI Study Platform
Handles extraction of text from PDF and DOCX files and generates AI-powered summaries
"""

import os
import re
import logging
from typing import Optional, Dict, List, Tuple
from io import BytesIO

# Document processing imports
import PyPDF2
from docx import Document

# AI/ML imports for summarization
try:
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    logging.warning("Transformers not available, falling back to extractive summarization")

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.frequency import FreqDist
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False
    logging.warning("NLTK not available, using basic summarization")

# File validation
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Main class for processing documents and generating summaries"""
    
    def __init__(self):
        self.summarizer = None
        self._initialize_summarizer()
    
    def _initialize_summarizer(self):
        """Initialize the AI summarizer if transformers is available"""
        if TRANSFORMERS_AVAILABLE:
            try:
                # Use a lightweight model for summarization
                self.summarizer = pipeline(
                    "summarization",
                    model="facebook/bart-large-cnn",
                    tokenizer="facebook/bart-large-cnn",
                    device=-1  # Use CPU
                )
                logger.info("BART summarization model loaded successfully")
            except Exception as e:
                logger.warning(f"Failed to load BART model: {e}")
                try:
                    # Fallback to smaller model
                    self.summarizer = pipeline(
                        "summarization",
                        model="sshleifer/distilbart-cnn-12-6",
                        device=-1
                    )
                    logger.info("DistilBART summarization model loaded successfully")
                except Exception as e2:
                    logger.warning(f"Failed to load DistilBART model: {e2}")
                    self.summarizer = None
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """Validate uploaded file type and content"""
        try:
            # Check file extension
            file_ext = filename.lower().split('.')[-1]
            if file_ext not in ['pdf', 'docx']:
                return False, "Only PDF and DOCX files are supported"
            
            # Check file size (max 10MB)
            if len(file_content) > 10 * 1024 * 1024:
                return False, "File size must be less than 10MB"
            
            # Basic magic number validation if available
            if MAGIC_AVAILABLE:
                file_type = magic.from_buffer(file_content, mime=True)
                if file_ext == 'pdf' and 'pdf' not in file_type:
                    return False, "Invalid PDF file"
                elif file_ext == 'docx' and 'document' not in file_type and 'zip' not in file_type:
                    return False, "Invalid DOCX file"
            
            return True, "File is valid"
            
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return False, f"File validation failed: {str(e)}"
    
    def extract_text_from_pdf(self, file_content: bytes) -> Tuple[bool, str, str]:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                return False, "", "PDF file appears to be empty"
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text.strip())
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                return False, "", "No readable text found in PDF"
            
            # Join all pages and clean up
            full_text = "\\n\\n".join(text_content)
            cleaned_text = self._clean_extracted_text(full_text)
            
            return True, cleaned_text, f"Successfully extracted text from {len(pdf_reader.pages)} pages"
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return False, "", f"Failed to process PDF: {str(e)}"
    
    def extract_text_from_docx(self, file_content: bytes) -> Tuple[bool, str, str]:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            if len(doc.paragraphs) == 0:
                return False, "", "DOCX file appears to be empty"
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            if not text_content:
                return False, "", "No readable text found in DOCX"
            
            # Join all content and clean up
            full_text = "\\n\\n".join(text_content)
            cleaned_text = self._clean_extracted_text(full_text)
            
            return True, cleaned_text, f"Successfully extracted text from document with {len(doc.paragraphs)} paragraphs"
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return False, "", f"Failed to process DOCX: {str(e)}"
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\\s+', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\\b\\d+\\b\\s*$', '', text, flags=re.MULTILINE)
        
        # Remove excessive line breaks
        text = re.sub(r'\\n{3,}', '\\n\\n', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\\w\\s.,;:!?\\(\\)\\[\\]\\-"\'\\n]', ' ', text)
        
        return text.strip()
    
    def generate_summary(self, text: str, max_length: int = 300) -> Tuple[bool, str, Dict]:
        """Generate AI-powered summary of the text"""
        try:
            if not text.strip():
                return False, "", {"error": "No text content to summarize"}
            
            original_word_count = len(text.split())
            
            # Split into chunks if text is too long
            chunks = self._split_text_into_chunks(text)
            
            # Try AI summarization first
            if self.summarizer and TRANSFORMERS_AVAILABLE:
                success, summary, metadata = self._generate_ai_summary(chunks, max_length)
            else:
                success, summary, metadata = self._generate_extractive_summary(text, max_length)
            
            # Final validation: ensure we have proper compression
            if success and summary:
                summary_word_count = len(summary.split())
                compression_ratio = (summary_word_count / original_word_count) * 100
                
                # If compression is poor (>90%), try more aggressive summarization
                if compression_ratio > 90 and original_word_count > 100:
                    logger.warning(f"Poor compression ratio: {compression_ratio:.1f}%, attempting more aggressive summarization")
                    
                    # Try extractive as fallback with aggressive settings
                    if self.summarizer and TRANSFORMERS_AVAILABLE:
                        # Already tried AI, try extractive fallback
                        success, summary, metadata = self._generate_extractive_summary(text, max_length)
                    else:
                        # Force more aggressive extractive summarization
                        sentences = re.split(r'[.!?]+', text)
                        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
                        if len(sentences) > 3:
                            # Take only top 2-3 sentences
                            num_sentences = max(2, min(3, len(sentences) // 5))
                            summary = '. '.join(sentences[:num_sentences]) + '.'
                            summary_word_count = len(summary.split())
                            metadata = {
                                "method": "Aggressive Basic Extractive",
                                "original_words": original_word_count,
                                "summary_words": summary_word_count,
                                "compression_ratio": round((summary_word_count / original_word_count) * 100, 1),
                                "note": "Applied aggressive compression due to poor initial ratio"
                            }
                
                # Update compression ratio in metadata
                final_compression_ratio = (len(summary.split()) / original_word_count) * 100
                metadata['final_compression_ratio'] = round(final_compression_ratio, 1)
                
                logger.info(f"📊 Summarization completed: {original_word_count} → {len(summary.split())} words ({final_compression_ratio:.1f}% compression)")
                
                return success, summary, metadata
            
            return success, summary, metadata
                
        except Exception as e:
            logger.error(f"Summarization error: {e}")
            return False, "", {"error": f"Failed to generate summary: {str(e)}"}
    
    def _split_text_into_chunks(self, text: str, max_chunk_size: int = 1000) -> List[str]:
        """Split text into manageable chunks for processing"""
        if NLTK_AVAILABLE:
            sentences = sent_tokenize(text)
        else:
            # Basic sentence splitting
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            if current_length + sentence_length <= max_chunk_size:
                current_chunk.append(sentence)
                current_length += sentence_length
            else:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_length = sentence_length
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def _generate_ai_summary(self, chunks: List[str], max_length: int) -> Tuple[bool, str, Dict]:
        """Generate summary using transformer models"""
        try:
            # Combine all chunks into one text for better context
            combined_text = ' '.join(chunks)
            original_word_count = len(combined_text.split())
            
            # If text is very short, don't summarize
            if original_word_count < 100:
                return True, combined_text, {
                    "method": "No summarization needed (too short)",
                    "original_length": original_word_count,
                    "summary_length": original_word_count
                }
            
            # Calculate appropriate summary length (aim for 25-40% of original)
            target_length = max(50, min(max_length, original_word_count // 3))
            min_length = max(30, target_length // 4)
            
            # Ensure we're actually compressing
            if target_length >= original_word_count * 0.8:
                target_length = max(50, int(original_word_count * 0.4))
                min_length = max(30, target_length // 3)
            
            # Generate summary
            result = self.summarizer(
                combined_text,
                max_length=target_length,
                min_length=min_length,
                do_sample=False,
                truncation=True,
                clean_up_tokenization_spaces=True
            )
            
            if result and len(result) > 0:
                summary = result[0]['summary_text'].strip()
                summary_word_count = len(summary.split())
                
                # Validate compression - if summary is too long, try again with stricter limits
                if summary_word_count >= original_word_count * 0.9:
                    logger.warning(f"Summary too long ({summary_word_count} vs {original_word_count} words), retrying with stricter limits")
                    
                    # Much more aggressive compression
                    strict_target = max(30, int(original_word_count * 0.25))
                    strict_min = max(20, strict_target // 3)
                    
                    retry_result = self.summarizer(
                        combined_text,
                        max_length=strict_target,
                        min_length=strict_min,
                        do_sample=False,
                        truncation=True,
                        clean_up_tokenization_spaces=True
                    )
                    
                    if retry_result and len(retry_result) > 0:
                        summary = retry_result[0]['summary_text'].strip()
                        summary_word_count = len(summary.split())
                
                metadata = {
                    "method": "AI (Transformer)",
                    "model": "BART",
                    "original_length": original_word_count,
                    "summary_length": summary_word_count,
                    "compression_ratio": round((summary_word_count / original_word_count) * 100, 1),
                    "target_length": target_length
                }
                
                return True, summary, metadata
            else:
                return False, "", {"error": "No summary generated from text"}
                
        except Exception as e:
            logger.error(f"AI summarization error: {e}")
            # Fallback to extractive summarization
            return self._generate_extractive_summary(' '.join(chunks), max_length)
    
    def _generate_extractive_summary(self, text: str, max_length: int) -> Tuple[bool, str, Dict]:
        """Generate summary using extractive methods (fallback)"""
        try:
            if NLTK_AVAILABLE:
                return self._nltk_extractive_summary(text, max_length)
            else:
                return self._basic_extractive_summary(text, max_length)
        except Exception as e:
            logger.error(f"Extractive summarization error: {e}")
            return False, "", {"error": f"Extractive summarization failed: {str(e)}"}
    
    def _nltk_extractive_summary(self, text: str, max_length: int) -> Tuple[bool, str, Dict]:
        """NLTK-based extractive summarization"""
        try:
            # Download required NLTK data if not present
            try:
                stopwords.words('english')
            except LookupError:
                nltk.download('stopwords')
            
            try:
                sent_tokenize('test')
            except LookupError:
                nltk.download('punkt')
            
            sentences = sent_tokenize(text)
            original_word_count = len(text.split())
            
            # If text is very short, don't summarize
            if len(sentences) <= 3 or original_word_count < 100:
                return True, text, {
                    "method": "No summarization needed (too short)", 
                    "sentences": len(sentences),
                    "original_words": original_word_count,
                    "summary_words": original_word_count
                }
            
            # Remove stopwords and calculate word frequency
            stop_words = set(stopwords.words('english'))
            words = word_tokenize(text.lower())
            words = [word for word in words if word.isalnum() and word not in stop_words]
            
            freq_dist = FreqDist(words)
            
            # Score sentences based on word frequency
            sentence_scores = {}
            for i, sentence in enumerate(sentences):
                sentence_words = word_tokenize(sentence.lower())
                score = 0
                word_count = 0
                
                for word in sentence_words:
                    if word in freq_dist:
                        score += freq_dist[word]
                        word_count += 1
                
                if word_count > 0:
                    sentence_scores[i] = score / word_count
                    # Boost early sentences
                    if i < len(sentences) * 0.3:
                        sentence_scores[i] *= 1.2
            
            # Calculate how many sentences to keep for proper compression
            # Aim for 30-50% compression
            target_sentences = max(2, min(len(sentences) // 2, len(sentences) // 3))
            
            # Ensure we don't select too many sentences
            if target_sentences > len(sentences) * 0.6:
                target_sentences = max(2, int(len(sentences) * 0.4))
            
            # Select top sentences
            top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:target_sentences]
            
            # Sort by original order
            top_sentences.sort(key=lambda x: x[0])
            
            summary_sentences = [sentences[i] for i, _ in top_sentences]
            summary = ' '.join(summary_sentences)
            summary_word_count = len(summary.split())
            
            # Validate compression - if summary is still too long, reduce further
            if summary_word_count >= original_word_count * 0.8:
                # More aggressive reduction
                aggressive_target = max(2, target_sentences // 2)
                top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:aggressive_target]
                top_sentences.sort(key=lambda x: x[0])
                summary_sentences = [sentences[i] for i, _ in top_sentences]
                summary = ' '.join(summary_sentences)
                summary_word_count = len(summary.split())
            
            metadata = {
                "method": "NLTK Extractive",
                "original_sentences": len(sentences),
                "summary_sentences": len(summary_sentences),
                "original_words": original_word_count,
                "summary_words": summary_word_count,
                "compression_ratio": round((summary_word_count / original_word_count) * 100, 1)
            }
            
            return True, summary, metadata
            
        except Exception as e:
            logger.error(f"NLTK summarization error: {e}")
            return self._basic_extractive_summary(text, max_length)
    
    def _basic_extractive_summary(self, text: str, max_length: int) -> Tuple[bool, str, Dict]:
        """Basic extractive summarization without external libraries"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        original_word_count = len(text.split())
        
        # If text is very short, don't summarize
        if len(sentences) <= 3 or original_word_count < 100:
            return True, text, {
                "method": "No summarization needed (too short)", 
                "sentences": len(sentences),
                "original_words": original_word_count,
                "summary_words": original_word_count
            }
        
        # Simple scoring based on sentence length and position
        scored_sentences = []
        for i, sentence in enumerate(sentences):
            score = len(sentence.split())  # Word count
            if i < len(sentences) * 0.3:  # First third gets bonus
                score *= 1.2
            if any(word in sentence.lower() for word in ['important', 'key', 'main', 'conclusion', 'summary']):
                score *= 1.1
            scored_sentences.append((sentence, score, i))
        
        # Calculate how many sentences to keep for proper compression
        # Aim for 30-50% compression
        target_sentences = max(2, min(len(sentences) // 2, len(sentences) // 3))
        
        # Ensure we don't select too many sentences
        if target_sentences > len(sentences) * 0.6:
            target_sentences = max(2, int(len(sentences) * 0.4))
        
        # Sort by score and take top sentences
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        top_sentences = scored_sentences[:target_sentences]
        
        # Sort back by original order
        top_sentences.sort(key=lambda x: x[2])
        
        summary = ' '.join([sentence for sentence, _, _ in top_sentences])
        summary_word_count = len(summary.split())
        
        # Validate compression - if summary is still too long, reduce further
        if summary_word_count >= original_word_count * 0.8:
            # More aggressive reduction
            aggressive_target = max(2, target_sentences // 2)
            top_sentences = scored_sentences[:aggressive_target]
            top_sentences.sort(key=lambda x: x[2])
            summary = ' '.join([sentence for sentence, _, _ in top_sentences])
            summary_word_count = len(summary.split())
        
        metadata = {
            "method": "Basic Extractive",
            "original_sentences": len(sentences),
            "summary_sentences": len(top_sentences),
            "original_words": original_word_count,
            "summary_words": summary_word_count,
            "compression_ratio": round((summary_word_count / original_word_count) * 100, 1)
        }
        
        return True, summary, metadata

# Global instance
document_processor = DocumentProcessor()

# Convenience functions
def process_document(file_content: bytes, filename: str) -> Dict:
    """Main function to process a document and return summary"""
    processor = document_processor
    
    logger.info(f"🔄 Processing document: {filename} ({len(file_content)} bytes)")
    
    # Validate file
    is_valid, validation_message = processor.validate_file(file_content, filename)
    logger.info(f"📋 File validation: {is_valid}, Message: {validation_message}")
    
    if not is_valid:
        return {
            "success": False,
            "error": validation_message,
            "extracted_text": "",
            "summary": "",
            "metadata": {}
        }
    
    # Extract text based on file type
    file_ext = filename.lower().split('.')[-1]
    logger.info(f"📄 Extracting text from {file_ext.upper()} file")
    
    if file_ext == 'pdf':
        success, extracted_text, message = processor.extract_text_from_pdf(file_content)
    elif file_ext == 'docx':
        success, extracted_text, message = processor.extract_text_from_docx(file_content)
    else:
        return {
            "success": False,
            "error": "Unsupported file format",
            "extracted_text": "",
            "summary": "",
            "metadata": {}
        }
    
    if not success:
        logger.error(f"❌ Text extraction failed: {message}")
        return {
            "success": False,
            "error": message,
            "extracted_text": "",
            "summary": "",
            "metadata": {}
        }
    
    logger.info(f"✅ Text extracted successfully. Length: {len(extracted_text)} characters")
    
    # Generate summary
    logger.info(f"🤖 Generating summary...")
    summary_success, summary, summary_metadata = processor.generate_summary(extracted_text)
    
    if not summary_success:
        logger.error(f"❌ Summary generation failed: {summary_metadata.get('error', 'Unknown error')}")
        return {
            "success": False,
            "error": summary_metadata.get("error", "Failed to generate summary"),
            "extracted_text": extracted_text,
            "summary": "",
            "metadata": summary_metadata
        }
    
    logger.info(f"✅ Document processing completed successfully!")
    logger.info(f"📊 Summary length: {len(summary)} characters")
    
    return {
        "success": True,
        "message": message,
        "extracted_text": extracted_text,
        "summary": summary,
        "metadata": {
            "extraction_info": message,
            "summary_info": summary_metadata,
            "file_size": len(file_content),
            "file_type": file_ext
        }
    }
